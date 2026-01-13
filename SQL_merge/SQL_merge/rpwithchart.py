import os
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')

# ==================== CELL FORMATTING ====================
def set_cell_bg(cell, fill_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)

# Text direction cho cell (dùng cho header và body)
def set_cell_text_direction(cell, direction="lrTb"):
    """
    direction options:
    - "btLr": vertical text, rotated 90° counterclockwise
    - "tbRl": vertical text, rotated 90° clockwise (thường dùng)
    - "lrTb": normal horizontal text
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement("w:textDirection")
    textDirection.set(qn("w:val"), direction)
    tcPr.append(textDirection)

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_cm * 567)))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def format_cell(cell, bold=False, font_color=None):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(12)
            run.bold = bold
            if font_color:
                run.font.color.rgb = font_color

def set_table_borders(table):
    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

def set_column_width(column, width_cm):
    """Set chiều rộng cột (width_cm: chiều rộng tính bằng cm)"""
    for cell in column.cells:
        cell.width = Inches(width_cm / 2.54)  # Convert cm to inches

def replace_placeholder_text(doc, placeholder, replacement):
    for p in doc.paragraphs:
        for run in p.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

# ==================== khúc này để vẽ chart ====================
def create_pie_chart(df, title, output_image, label_col_idx=0, value_col_idx=1, top_n=10):
    try:
        df_chart = df.head(top_n).copy()
        labels = df_chart.iloc[:, label_col_idx].astype(str).tolist()
        values = pd.to_numeric(df_chart.iloc[:, value_col_idx], errors='coerce').fillna(0).tolist()
        
        valid_data = [(label, value) for label, value in zip(labels, values) 
                      if value > 0 and label.strip() != '' and label.lower() not in ['nan', 'none']]
        
        if not valid_data:
            print(f"   ⚠️ No valid data for chart: {title}")
            return False
        
        labels, values = zip(*valid_data)
        
        fig, ax = plt.subplots(figsize=(10, 8), facecolor='white')
        colors = ['#5B9BD5', '#ED7D31', '#A5A5A5', '#FFC000', '#70AD47', 
                  '#4472C4', '#C55A11', '#7030A0', '#44546A', '#264478']
        
        wedges, texts = ax.pie(values, labels=None, startangle=90,
                               colors=colors[:len(values)], explode=[0.02] * len(values))
        
        ax.set_title(title, fontsize=16, fontweight='bold', pad=30, color='#333333')
        
        num_cols = 3 if len(labels) > 6 else (2 if len(labels) > 3 else 1)
        legend = ax.legend(labels, loc='upper center', bbox_to_anchor=(0.5, -0.05),
                          ncol=num_cols, frameon=False, fontsize=11)
        
        try:
            for i, wedge in enumerate(wedges):
                if i < len(legend.get_patches()):
                    legend.get_patches()[i].set_facecolor(colors[i % len(colors)])
        except:
            pass
        
        plt.axis('equal')
        plt.tight_layout()
        plt.savefig(output_image, dpi=150, bbox_inches='tight', facecolor='white')
        plt.close()
        
        print(f"    ✅ Created chart: {output_image}")
        return True
    except Exception as e:
        print(f"   ❌ Error creating chart: {e}")
        plt.close()
        return False

# ==================== MAIN REPORT GENERATOR ====================
def generate_report(excel_file: str, template_file: str, output_file: str, mapping: dict, chart_mapping: dict = None):
    xls = pd.ExcelFile(excel_file)
    doc = Document(template_file)
    temp_images = []

    # Xử lý các placeholder bảng
    for placeholder, config in mapping.items():
        try:
            # Xử lý collect_date
            if placeholder == "<collect_date>":
                current_date = datetime.now().strftime("%m.%Y")
                replace_placeholder_text(doc, placeholder, current_date)
                print(f"✅ Replaced {placeholder} with {current_date}")
                continue

            if not config:
                continue

            sheet_name = config["sheet"]
            df = pd.read_excel(xls, sheet_name=sheet_name)
            
            # ========== TRANSPOSE LOGIC ==========
            if config.get("transpose", False):
                # Chọn cột trước khi transpose
                if "columns" in config and config["columns"]:
                    col_indices = config["columns"]
                    selected_cols = [df.columns[i] for i in col_indices if i < len(df.columns)]
                    df = df[selected_cols]
                
                # Lưu tên cột đầu tiên gốc (ví dụ: "volume_mount_point")
                original_first_col = df.columns[0]
                
                # Transpose: cột thành hàng, hàng thành cột
                df = df.T.reset_index()
                
                # Đặt tên cột: cột đầu dùng tên gốc, các cột sau lấy từ hàng đầu tiên
                if len(df.columns) > 1:
                    # Lấy giá trị hàng đầu làm tên cột (ví dụ: D:\, E:\)
                    new_columns = [original_first_col] + [str(val) for val in df.iloc[0, 1:].tolist()]
                    df.columns = new_columns
                    # Xóa hàng đầu vì đã dùng làm header
                    df = df.iloc[1:].reset_index(drop=True)
                
                # Loại bỏ các cột có tên chứa 'nan' hoặc rỗng
                df = df.loc[:, ~df.columns.str.lower().str.contains('nan', na=False)]
                df = df.loc[:, df.columns.str.strip() != '']
            # ====================================
            else:
                # Không transpose: chọn cột bình thường
                if "columns" in config and config["columns"]:
                    col_indices = config["columns"]
                    selected = [df.columns[i] for i in col_indices if i < len(df.columns)]
                    df = df[selected]

            # Giới hạn số hàng (áp dụng sau khi transpose hoặc không transpose)
            max_rows = config.get("max_rows", None)
            if max_rows and len(df) > max_rows:
                df = df.head(max_rows)

            # Tìm placeholder và chèn bảng
            for p in doc.paragraphs:
                if placeholder in p.text:
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.autofit = True
                    set_table_borders(table)
                    
                    hdr_cells = table.rows[0].cells
                    header_height = config.get("header_height", 1.8)
                    set_row_height(table.rows[0], header_height)
                    
                    # ========== TEXT DIRECTION CHO HEADER ==========
                    use_vertical_header = config.get("vertical_header", False)
                    for j, col in enumerate(df.columns):
                        hdr_cells[j].text = str(col)
                        set_cell_bg(hdr_cells[j], "0066CC")
                        format_cell(hdr_cells[j], bold=True, font_color=RGBColor(255, 255, 255))
                        if use_vertical_header:
                            set_cell_text_direction(hdr_cells[j], "tbRl")
                    # ==============================================

                    # ========== TEXT DIRECTION CHO BODY CELLS ==========
                    horizontal_columns = config.get("horizontal_columns", [])
                    vertical_body = config.get("vertical_body", False)
                    row_height = config.get("row_height", 1.8)
                    
                    for _, row in df.iterrows():
                        new_row = table.add_row()
                        row_cells = new_row.cells
                        set_row_height(new_row, row_height)
                        
                        for j, val in enumerate(row):
                            row_cells[j].text = str(val)
                            format_cell(row_cells[j])
                            
                            # Chỉ apply vertical nếu cột không nằm trong horizontal_columns
                            if vertical_body:
                                col_name = df.columns[j]
                                if col_name not in horizontal_columns:
                                    set_cell_text_direction(row_cells[j], "tbRl")
                    # ==================================================

                    p._element.getparent().replace(p._element, table._element)
                    
                    # ========== SET COLUMN WIDTH (nếu có config) ==========
                    if "column_widths" in config:
                        col_widths = config["column_widths"]
                        for col_idx, width_cm in enumerate(col_widths):
                            if col_idx < len(table.columns):
                                set_column_width(table.columns[col_idx], width_cm)
                    # =====================================================

            print(f"✅ Replaced {placeholder} with sheet '{sheet_name}' (rows={len(df)})")

        except Exception as e:
            print(f"⚠️ Could not process {placeholder}: {e}")

    # ========== chèn CHART PLACEHOLDERS ==========
    if chart_mapping:
        for placeholder, config in chart_mapping.items():
            try:
                sheet_name = config["sheet"]
                chart_title = config.get("title", sheet_name)
                df = pd.read_excel(xls, sheet_name=sheet_name)
                
                temp_image = f"temp_chart_{placeholder.strip('<>').replace('_', '')}.png"
                temp_images.append(temp_image)
                
                label_col = config.get("label_col", 0)
                value_col = config.get("value_col", 1)
                top_n = config.get("top_n", 10)
                
                if create_pie_chart(df, chart_title, temp_image, label_col, value_col, top_n):
                    for p in doc.paragraphs:
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, "")
                            run = p.add_run()
                            run.add_picture(temp_image, width=Inches(5.5))
                            print(f"✅ Inserted chart for {placeholder}")
                            break
            except Exception as e:
                print(f"⚠️ Could not create chart for {placeholder}: {e}")
    # =============================================

    # Lưu file
    try:
        doc.save(output_file)
        print(f"\n✅ Report generated: {output_file}")
    except PermissionError:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = os.path.splitext(output_file)[0]
        new_output = f"{base_name}_{timestamp}.docx"
        doc.save(new_output)
        print(f"\n⚠️ File đang mở. Đã lưu thành: {new_output}")
    
    # Cleanup temp images
    for temp_img in temp_images:
        try:
            if os.path.exists(temp_img):
                os.remove(temp_img)
        except:
            pass

# ==================== MAIN EXECUTION ====================
if __name__ == "__main__":
    template_folder = r"D:\INTERNSHIP\SQL_merge_260112\SQL_merge\SQL_merge\rptemplate"
    excel_folder = r"D:\INTERNSHIP\SQL_merge_260112\SQL_merge\SQL_merge\output"
    output_folder = r"D:\INTERNSHIP\SQL_merge_260112\SQL_merge\SQL_merge\reports"
    os.makedirs(output_folder, exist_ok=True)

    # ========== MAPPING CỦA CÁC BẢNG ==========
    mapping = {
        # TRANSPOSE: Chuyển cột thành hàng (tự động lấy header từ data)
        "<volume_info>": {
            "sheet": "Volume Info",
            "columns": [0, 1, 2, 3, 4, 5],
            "transpose": True
        },
        
        "<file_size>": {
            "sheet": "File Sizes and Space",
            "columns": [0, 1, 2, 3, 4, 5, 7],
            "max_rows": 50
        },
        
        # TEXT DIRECTION: cho cái sheet fileio
        "<fileio>": {
            "sheet": "IO Stats By File",
            "max_rows": 50,
            "vertical_header": True,
            "vertical_body": True,
            "horizontal_columns": ["Database Name", "Logical Name", "type_desc", "Physical Name", "file_id"],
            "header_height": 2.0,
            "row_height": 1.8,
            "column_widths": [2.5, 2.5, 1.2, 2.0, 10.0, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]  
        },
        
        "<conn_count>": {
            "sheet": "Connection Counts by IP Address",
            "max_rows": 50
        },
        "<cpu_usage>": {
            "sheet": "CPU Usage by Database",
            "columns": [0, 1, 3],
            "max_rows": 50
        },
        "<io_usage>": {
            "sheet": "IO Usage By Database",
            "columns": [0, 1, 3],
            "max_rows": 50
        },
        "<buffer_usage>": {
            "sheet": "Total Buffer Usage by Database",
            "columns": [0, 1, 3],
            "max_rows": 50
        },
        "<top_worker>": {
            "sheet": "Top Worker Time Queries",
            "columns": [0, 1, 2, 4],
            "max_rows": 50
        },
        "<missing_index>": {
            "sheet": "Missing Indexes",
            "columns": [2, 5, 6, 7, 9],
            "max_rows": 50
        },
        "<agent_job>": {
            "sheet": "SQL Server Agent Jobs",
            "columns": [0, 1, 2, 3, 4, 8, 9],
            "max_rows": 50
        },
        "<recent_bk>": {
            "sheet": "Recent Full Backups",
            "columns": [2, 3, 4, 5, 11],
            "max_rows": 50
        },
        "<collect_date>": {}
    }

    # ========== CHART MAPPING ==========
    chart_mapping = {
        "<cpu_usage_chart>": {
            "sheet": "CPU Usage by Database",
            "title": "Chart 1. CPU Usage by Database",
            "label_col": 1,
            "value_col": 3,
            "top_n": 10
        },
        "<io_usage_chart>": {
            "sheet": "IO Usage By Database",
            "title": "Chart 2. IO Usage By Database",
            "label_col": 1,
            "value_col": 3,
            "top_n": 10
        },
        "<buffer_usage_chart>": {
            "sheet": "Total Buffer Usage by Database",
            "title": "Chart 3. Total Buffer Usage by Database",
            "label_col": 1,
            "value_col": 3,
            "top_n": 10
        }
    }

    # ========== XỬ LÝ TẤT CẢ TEMPLATE ==========
    for template_file in os.listdir(template_folder):
        if not template_file.lower().endswith(".docx") or template_file.startswith("~$"):
            continue

        base_name = os.path.splitext(template_file)[0]
        keyword = None
        for part in base_name.split("_"):
            if part.startswith("INS"):
                keyword = part
                break

        if not keyword:
            print(f"⚠️ Không tìm thấy keyword 'INS...' trong {template_file}")
            continue

        excel_match = None
        for excel_file in os.listdir(excel_folder):
            if keyword in excel_file and excel_file.lower().endswith(".xlsx") and not excel_file.startswith("~$"):
                excel_match = excel_file
                break

        if not excel_match:
            print(f"⚠️ Không tìm thấy excel cho template {template_file} (keyword={keyword})")
            continue

        output_file = os.path.join(output_folder, f"{template_file}")

        print(f"\n{'='*60}")
        print(f" Processing: {template_file}")
        print(f" Excel: {excel_match}")
        print(f"{'='*60}")

        generate_report(
            excel_file=os.path.join(excel_folder, excel_match),
            template_file=os.path.join(template_folder, template_file),
            output_file=output_file,
            mapping=mapping,
            chart_mapping=chart_mapping
        )