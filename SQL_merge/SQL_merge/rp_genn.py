import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_bg(cell, fill_color: str):
    """
    Set background color for a cell (fill_color dạng hex 'RRGGBB').
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_color)
    tcPr.append(shd)

def format_cell(cell, bold=False, font_color=None):
    """
    Format text trong cell: font Cambria, size 12, optional bold + font color.
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Cambria"
            run.font.size = Pt(12)
            run.bold = bold
            if font_color:
                run.font.color.rgb = font_color

def generate_report(excel_file: str, template_file: str, output_file: str,
                    mapping: dict):
    """
    Generate Word report by replacing placeholders in template with Excel sheet data.
    - Chỉ chọn cột theo index (0-based).
    - Xuất dữ liệu ra bảng Word, có style Cambria 12, header format.
    """
    # Load Excel
    xls = pd.ExcelFile(excel_file)

    # Load Word template
    doc = Document(template_file)

    # Replace placeholders
    for placeholder, config in mapping.items():
        try:
            sheet_name = config["sheet"]
            df = pd.read_excel(xls, sheet_name=sheet_name)

            # Lọc cột theo index
            if "columns" in config and config["columns"]:
                col_indices = config["columns"]
                selected = [df.columns[i] for i in col_indices if i < len(df.columns)]
                df = df[selected]

            # Giới hạn số dòng
            max_rows = config.get("max_rows", None)
            if max_rows and len(df) > max_rows:
                df = df.head(max_rows)

            # ---- Insert table in place of placeholder ----
            for p in doc.paragraphs:
                if placeholder in p.text:
                    # Tạo table từ DataFrame
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.autofit = True

                    # Header row
                    hdr_cells = table.rows[0].cells
                    for j, col in enumerate(df.columns):
                        hdr_cells[j].text = str(col)
                        set_cell_bg(hdr_cells[j], "0066CC")
                        format_cell(hdr_cells[j], bold=True, font_color=RGBColor(255, 255, 255))

                    # Data rows
                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, val in enumerate(row):
                            row_cells[j].text = str(val)
                            format_cell(row_cells[j])

                    # Xóa placeholder text
                    p.text = p.text.replace(placeholder, "")
                    # Chèn bảng ngay sau paragraph đó
                    p._element.addnext(table._element)

            print(f"✅ Replaced {placeholder} with sheet '{sheet_name}' (rows={len(df)})")

        except Exception as e:
            print(f"⚠️ Could not process {placeholder}: {e}")

    # Save new report
    doc.save(output_file)
    print(f"\n📄 Report generated: {output_file}")


if __name__ == "__main__":
    excel_file = r"D:\SQL_merge\output\105DBCF_HealthCheck_AllInfo.xlsx"
    template_file = r"D:\SQL_merge\rptemplate\test\SGC_SQL_HEALTHCHECK_INS105DCDBCF.docx"
    output_file = r"D:\SQL_merge\report105.docx"

    # Mapping placeholder -> sheet config
    mapping = {
                "<file_size>": {"sheet": "File Sizes and Space", "columns": [0, 1, 2, 3, 4, 5, 7], "max_rows": 50},
        "<fileio>": {"sheet": "IO Stats By File", "max_rows": 50},
        "<conn_count>": {"sheet": "Connection Counts by IP Address", "max_rows": 50},
        "<cpu_usage>": {"sheet": "CPU Usage by Database", "columns": [0, 1, 3], "max_rows": 50},
        "<io_usage>": {"sheet": "IO Usage By Database", "columns": [0, 1, 3], "max_rows": 50},
        "<buffer_usage>": {"sheet": "Total Buffer Usage by Database", "columns": [0, 1, 3], "max_rows": 50},
        "<top_worker>": {"sheet": "Top Worker Time Queries", "columns": [0, 1, 2, 4], "max_rows": 50},
        "<missing_index>": {"sheet": "Missing Indexes", "columns": [2, 5, 6, 7, 9], "max_rows": 50},
        "<agent_job>": {"sheet": "SQL Server Agent Jobs", "columns": [0, 1, 2, 3, 4, 8, 9], "max_rows": 50},
        "<recent_bk>": {"sheet": "Recent Full Backups", "columns": [2, 3, 4, 5, 11], "max_rows": 50},
        "<collect_date>": {},
    }

    generate_report(excel_file, template_file, output_file, mapping)
